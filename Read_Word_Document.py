import docx
import pandas as pd
import pyodbc


def get_column_name(key, df):
    col = df.filter(like=key).columns.tolist()
    return col[0] if col.__len__() > 0 else ""


def connect_microsoft_server_management(id_macro_ent, data_date):
    output = list()
    connection = pyodbc.connect('Driver={SQL Server}; Server=dirac\\dirac2012; Database=Tier2; Trusted_Connection=yes;')
    cursor = connection.cursor()
    test_sql = "SELECT * FROM Tier2.dat.macro WHERE id_macro_ent=" + id_macro_ent + " and Data_date=" \
               + "'" + data_date + "'"
    cursor.execute(test_sql)
    for row in cursor:
        print(row)
        output.append(row)
    return output


def save_xls(list_dfs, xls_path):
    with pd.ExcelWriter(xls_path) as writer:
        sheet_list = list()
        for n, df in enumerate(list_dfs):
            df.to_excel(writer, 'sheet%s' % n)
            sheet_list.append('sheet%s' % n)
        writer.save()
        return sheet_list


def confirm_econ_tmr(keys, string_i):
    for key in keys:
        if key in string_i:
            return True
    return False


def generate_detail(file, reference, prefix):
    details = list()
    index_name_list = list()
    doc = docx.Document(file)
    index_name = ""
    for i in doc.paragraphs:
        i = i.text.replace(";", "")
        if str(i).startswith('GC') and confirm_econ_tmr(reference.keys(), i):
            if "%" in i:
                index_name = i.split("%")[1].strip()
                i = i.split("%")[0].strip()
            index = int(i.split("{")[1].split("}")[0].strip())
            value = i.split("=")[1].replace("{", "").replace("}", "").replace("[", "").replace("]", "") \
                .replace("'", "").replace(",", ", ").replace(",", "").strip()
            if " " in value:
                value = value.split(" ")[-1]
            list_name = i.split("GC.")[1].split("{")[0].strip()
            reference[list_name].append([index, value])
            if list_name == list(reference.keys())[0]:
                index_name_list.append(index_name)
        else:
            pass
    for key in reference.keys():
        list_value = reference[key]
        list_value = pd.DataFrame(list_value, columns=['Index', key])
        list_value.set_index('Index', inplace=True)
        details.append(list_value)
    details = pd.concat(details, axis='columns')
    details['Index Name ' + prefix] = index_name_list
    details = details.sort_index()
    return details


def generate_econ(file):
    country_name = list()
    doc = docx.Document(file)
    for i in doc.paragraphs:
        i = i.text
        if str(i).startswith('GC'):
            index = int(i.split("{")[1].split("}")[0])
            country = i.split("'")[1]
            country_name.append([index, country])
        else:
            pass
    countries = pd.DataFrame(country_name, columns=['Index', 'Country'])
    countries.set_index('Index', inplace=True)
    countries = countries.sort_index()
    return countries


def read_filtered_result(file_path, sheet_list):
    xlsx = pd.ExcelFile(file_path)
    id_list = list()
    for sheet in sheet_list:
        df = pd.read_excel(xlsx, sheet)
        target_column_var = get_column_name('VAR', df)
        for i in df[target_column_var]:
            id_list.append(i)
    print(id_list)
    return id_list


def main():
    econ_name = generate_econ('ECON.docx')
    prefix = 'tmr'
    tmr_var = list()
    tmr_source = list()
    tmr_key_date = list()
    tmr_adj_date = list()
    reference_tmr = {'TMR_VAR': tmr_var, 'TMR_SOURCE': tmr_source, 'TMR_KEYDATE': tmr_key_date,
                     'TMR_ADJDATE': tmr_adj_date}
    econ_tmr = generate_detail('TMR.docx', reference_tmr, prefix)
    prefix = 'stkidx'
    stkidx_var = list()
    stkidx_source = list()
    stkidx_keydate = list()
    stkidx_currency = list()
    reference_stkidx = {'STKIDX_VAR': stkidx_var, 'STKIDX_SOURCE': stkidx_source, 'STKIDX_KEYDATE': stkidx_keydate,
                        'STKIDX_CURRENCY': stkidx_currency}
    econ_stkidx = generate_detail('STKIDX.docx', reference_stkidx, prefix)
    prefix = 'rfr'
    rfr_var = list()
    rfr_source = list()
    rfr_keydate = list()
    reference_rfr = {'RFR_VAR': rfr_var, 'RFR_SOURCE': rfr_source, 'RFR_KEYDATE': rfr_keydate}
    econ_rfr = generate_detail('RFR.docx', reference_rfr, prefix)
    prefix = 'fxr'
    fxr_var = list()
    fxr_source = list()
    fxr_keydate = list()
    reference_fxr = {'FXR_VAR': fxr_var, 'FXR_SOURCE': fxr_source, 'FXR_KEYDATE': fxr_keydate}
    econ_fxr = generate_detail('FXR.docx', reference_fxr, prefix)
    ele_list = [econ_tmr, econ_stkidx, econ_rfr, econ_fxr]
    sheet_list = list()
    sheet_list_filtered = list()
    for ele in ele_list:
        details = pd.concat([econ_name, ele], axis='columns', sort=True)
        details = pd.DataFrame(details)
        target_column_source = get_column_name('SOURCE', details)
        target_column_var = get_column_name('VAR', details)
        target_column_keydate = get_column_name('KEYDATE', details)
        target_column_adjdate = get_column_name('ADJDATE', details)
        target_column_currency = get_column_name('CURRENCY', details)
        numeric_column_list = [target_column_var, target_column_keydate, target_column_adjdate, target_column_currency]
        condition_source = details[target_column_source] == "DS"
        filtered = details[condition_source][target_column_source].dropna()
        filtered = details.loc[filtered.index]
        for i in numeric_column_list:
            if i in filtered.columns:
                filtered[i] = pd.to_numeric(filtered[i])
        sheet_list.append(details)
        sheet_list_filtered.append(filtered)
    sheet_list_result = save_xls(sheet_list, "Result.xlsx")
    sheet_list_filtered = save_xls(sheet_list_filtered, "Result_filtered.xlsx")
    read_filtered_result('Result_filtered.xlsx', sheet_list_filtered)


if __name__ == '__main__':
    main()
